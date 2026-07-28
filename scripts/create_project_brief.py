"""Generate the MotionBridge project brief as a polished Word document."""

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "MotionBridge_Project_Brief.docx"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
PALE_BLUE = "E8EEF5"
PALE_GREEN = "E2F0D9"
PALE_AMBER = "FFF2CC"
PALE_RED = "FCE4D6"
LIGHT_GREY = "F2F2F2"
TEXT = RGBColor(35, 43, 53)


def set_cell_shading(cell, fill):
    props = cell._tc.get_or_add_tcPr()
    shading = props.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        props.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_width(cell, twips):
    props = cell._tc.get_or_add_tcPr()
    width = props.find(qn("w:tcW"))
    if width is None:
        width = OxmlElement("w:tcW")
        props.append(width)
    width.set(qn("w:w"), str(twips))
    width.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    props = table._tbl.tblPr
    table_width = props.find(qn("w:tblW"))
    if table_width is None:
        table_width = OxmlElement("w:tblW")
        props.append(table_width)
    table_width.set(qn("w:w"), str(sum(widths)))
    table_width.set(qn("w:type"), "dxa")
    table_indent = props.find(qn("w:tblInd"))
    if table_indent is None:
        table_indent = OxmlElement("w:tblInd")
        props.append(table_indent)
    table_indent.set(qn("w:w"), "120")
    table_indent.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            set_cell_width(cell, width)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            margins = cell._tc.get_or_add_tcPr()
            cell_mar = margins.find(qn("w:tcMar"))
            if cell_mar is None:
                cell_mar = OxmlElement("w:tcMar")
                margins.append(cell_mar)
            for side, value in (("top", 80), ("bottom", 80), ("start", 120), ("end", 120)):
                node = OxmlElement(f"w:{side}")
                node.set(qn("w:w"), str(value))
                node.set(qn("w:type"), "dxa")
                cell_mar.append(node)
    first_row_props = table.rows[0]._tr.get_or_add_trPr()
    repeat = first_row_props.find(qn("w:tblHeader"))
    if repeat is None:
        repeat = OxmlElement("w:tblHeader")
        first_row_props.append(repeat)
    repeat.set(qn("w:val"), "true")


def add_table(doc, headers, rows, widths, status_colors=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, heading in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, PALE_BLUE)
        paragraph = cell.paragraphs[0]
        run = paragraph.add_run(heading)
        run.bold = True
        run.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    for row_data in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row_data):
            cells[idx].text = str(value)
            if status_colors and idx == status_colors[0]:
                fill = status_colors[1].get(str(value))
                if fill:
                    set_cell_shading(cells[idx], fill)
    set_table_geometry(table, widths)
    doc.add_paragraph()
    return table


def add_bullet(doc, text, level=0):
    paragraph = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    paragraph.add_run(text)
    return paragraph


def add_number(doc, text):
    paragraph = doc.add_paragraph(style="List Number")
    paragraph.add_run(text)
    return paragraph


def add_callout(doc, title, text, fill=PALE_BLUE):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(f"{title}\n")
    run.bold = True
    run.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    paragraph.add_run(text)
    set_table_geometry(table, [9360])
    doc.add_paragraph()


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("MotionBridge  |  ")
    run.font.size = Pt(9)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    paragraph._p.append(field)


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)
    section.different_first_page_header_footer = True

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = TEXT
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15
    for name, size, color, before, after in (
        ("Title", 31, DARK_BLUE, 0, 12),
        ("Subtitle", 15, BLUE, 0, 18),
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 11.5, DARK_BLUE, 10, 5),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
    styles["Title"].font.bold = True
    styles["Heading 1"].font.bold = True
    styles["Heading 2"].font.bold = True
    styles["Heading 3"].font.bold = True
    for list_name in ("List Bullet", "List Bullet 2", "List Number"):
        styles[list_name].font.name = "Calibri"
        styles[list_name].font.size = Pt(10.5)
        styles[list_name].paragraph_format.space_after = Pt(4)

    for sec in doc.sections:
        header = sec.header
        header.is_linked_to_previous = False
        p = header.paragraphs[0]
        p.text = "MOTIONBRIDGE  /  PROJECT BRIEF"
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.runs[0].font.size = Pt(8)
        p.runs[0].font.bold = True
        p.runs[0].font.color.rgb = RGBColor.from_string(BLUE)
        footer = sec.footer
        footer.is_linked_to_previous = False
        add_page_number(footer.paragraphs[0])


def create_document():
    doc = Document()
    configure_document(doc)
    props = doc.core_properties
    props.title = "MotionBridge — Project Brief"
    props.subject = "C++20 industrial motion-control portfolio project"
    props.author = "MotionBridge Project"
    props.keywords = "C++20, PLC, OPC UA, motion control, ROS 2, EtherCAT"

    # Cover
    doc.add_paragraph("\n\n")
    eyebrow = doc.add_paragraph()
    eyebrow.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = eyebrow.add_run("INDUSTRIAL CONTROL • C++20 • SIMULATION")
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor.from_string(BLUE)
    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("MotionBridge")
    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("Vendor-Independent Industrial Motion Control Stack")
    lead = doc.add_paragraph()
    lead.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lead.paragraph_format.left_indent = Inches(0.55)
    lead.paragraph_format.right_indent = Inches(0.55)
    lead.add_run(
        "A practical repository showing how a PLC can supervise a deterministic-style "
        "C++ motion controller while networking, robotics, and fieldbus technologies "
        "remain replaceable adapters."
    )
    doc.add_paragraph()
    metrics = doc.add_table(rows=1, cols=3)
    values = (("1 kHz", "control loop"), ("50 Hz", "PLC worker"), ("12 / 12", "tests passing"))
    for cell, (value, label) in zip(metrics.rows[0].cells, values):
        set_cell_shading(cell, PALE_BLUE)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(value + "\n")
        r.bold = True
        r.font.size = Pt(17)
        r.font.color.rgb = RGBColor.from_string(DARK_BLUE)
        p.add_run(label).font.size = Pt(9)
    set_table_geometry(metrics, [3120, 3120, 3120])
    doc.add_paragraph("\n")
    status = doc.add_paragraph()
    status.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = status.add_run("CURRENT STATUS  •  LOCAL CONTROLLER + MOCK PLC + OPC UA CLIENT BOUNDARY")
    r.bold = True
    r.font.size = Pt(9.5)
    r.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    date = doc.add_paragraph("Project brief • 29 July 2026")
    date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date.runs[0].italic = True
    date.runs[0].font.color.rgb = RGBColor(90, 100, 110)
    doc.add_page_break()

    doc.add_heading("1. The project in one page", level=1)
    doc.add_heading("Objective", level=2)
    doc.add_paragraph(
        "Build a modular single-axis motion-control platform in modern C++ that can be "
        "demonstrated without hardware today, then connected to Siemens PLC, ROS 2, "
        "OPC UA, and EtherCAT/TwinCAT without rewriting the control kernel."
    )
    doc.add_heading("Problem being addressed", level=2)
    add_bullet(doc, "Industrial systems mix slow supervisory decisions with fast control work.")
    add_bullet(doc, "Network calls are unpredictable and must not block the high-frequency loop.")
    add_bullet(doc, "Real hardware and vendor software are expensive or unavailable during development.")
    add_bullet(doc, "Safety behavior must be explicit: commands are validated, faults latch, and reset is deliberate.")
    doc.add_heading("The core design decision", level=2)
    add_callout(
        doc,
        "Separation of responsibility",
        "PLC = machine supervision and permission. C++ = trajectory and closed-loop motion. "
        "OPC UA = supervisory networking. EtherCAT = cyclic drive data. ROS 2 = higher-level goals and diagnostics.",
    )
    doc.add_heading("What works now", level=2)
    add_table(
        doc,
        ["Capability", "Status", "What proves it"],
        [
            ("C++20 control kernel", "Complete", "State machine, trajectory, PID, plant, fixed-period loop"),
            ("Visual simulation", "Complete", "CSV telemetry and SVG plots"),
            ("PLC behavior", "Complete", "Mock PLC, control/status words, heartbeat, watchdog"),
            ("Thread separation", "Complete", "50 Hz worker + newest-value mailboxes + 1 kHz loop"),
            ("OPC UA boundary", "Complete", "Optional open62541 adapter and connection probe"),
            ("Live Siemens link", "Next", "TIA/PLCSIM server configuration and exact NodeIds"),
            ("ROS 2 and EtherCAT", "Planned", "Adapters after the Siemens integration"),
        ],
        [2700, 1300, 5360],
        (1, {"Complete": PALE_GREEN, "Next": PALE_AMBER, "Planned": LIGHT_GREY}),
    )
    add_callout(
        doc,
        "Honest scope",
        "This is a functional engineering simulation and portfolio project. It is not a "
        "safety-certified controller and Windows measurements do not claim hard real-time behavior.",
        PALE_AMBER,
    )

    doc.add_page_break()
    doc.add_heading("2. Architecture and technology roles", level=1)
    add_table(
        doc,
        ["Layer", "Rate", "Responsibility", "Current implementation"],
        [
            ("ROS 2", "10–100 Hz", "Goals, telemetry, visualization", "Planned adapter"),
            ("PLC / OPC UA", "50 Hz", "Enable, start, stop, reset, heartbeat", "Mock PLC + real client boundary"),
            ("C++ controller", "1,000 Hz", "State, trajectory, PID, limits, diagnostics", "Implemented"),
            ("Fieldbus / drive", "500–1,000 Hz", "Process data and actuator feedback", "Simulated servo plant"),
        ],
        [1500, 1200, 3100, 3560],
    )
    doc.add_heading("Signal flow", level=2)
    add_table(
        doc,
        ["1. Input", "2. Transfer", "3. Control", "4. Plant", "5. Feedback"],
        [
            (
                "PLC / future ROS goal",
                "50 Hz worker + mailbox",
                "State → trajectory → PID",
                "Simulated motor / future drive",
                "Position, state, fault, timing",
            )
        ],
        [1740, 1800, 2040, 1920, 1860],
    )
    doc.add_paragraph(
        "The control loop never waits for OPC UA. The communication worker performs network "
        "reads and writes, then exchanges only the newest command and status through bounded "
        "mailboxes. This keeps network delays outside the controller’s timing path."
    )
    doc.add_heading("Technologies and why they are used", level=2)
    add_table(
        doc,
        ["Technology", "Purpose in MotionBridge"],
        [
            ("C++20", "Type-safe core, RAII, threads, chrono timing, portable interfaces"),
            ("CMake + CTest", "Repeatable builds and automated verification"),
            ("open62541", "Optional real OPC UA client for Siemens or other vendors"),
            ("TIA Portal / PLCSIM Advanced", "Next supervisory PLC and OPC UA server"),
            ("CSV + SVG", "Visible proof of reference tracking and controller response"),
            ("ROS 2 / ros2_control", "Planned robotics-facing hardware abstraction"),
            ("EtherCAT / CiA 402", "Planned cyclic process-data and drive-state backend"),
            ("Git", "Traceable milestones and portfolio delivery"),
        ],
        [2300, 7060],
    )

    doc.add_page_break()
    doc.add_heading("3. Inputs, processing, and outputs", level=1)
    doc.add_heading("Supervisory inputs", level=2)
    add_table(
        doc,
        ["Input", "Meaning", "Typical rule"],
        [
            ("Enable", "Permit the axis to become ready", "Ignored when a fault or E-stop is active"),
            ("Start", "Execute the accepted target", "Requires enabled/ready state"),
            ("Stop", "Request controlled stopping", "Removes active motion request"),
            ("Reset fault", "Clear a latched recoverable fault", "Accepted only after the cause is gone"),
            ("Emergency stop", "Force emergency-stop state", "Highest-priority supervisory command"),
            ("Target position", "Requested axis position [rad]", "Must be finite"),
            ("Maximum velocity", "Trajectory speed limit [rad/s]", "Must be positive and finite"),
            ("Maximum acceleration", "Trajectory acceleration limit [rad/s²]", "Must be positive and finite"),
            ("PLC heartbeat", "Evidence of fresh communication", "Must keep changing within timeout"),
        ],
        [2200, 3900, 3260],
    )
    doc.add_heading("Processing pipeline", level=2)
    for text in (
        "Receive the newest supervisory command.",
        "Validate numerical values and check heartbeat freshness.",
        "Advance the industrial state machine.",
        "Generate a velocity-limited trapezoidal position reference.",
        "Calculate PID torque with saturation and anti-windup.",
        "Update the simulated servo plant and read its feedback.",
        "Publish status, faults, heartbeats, and timing measurements.",
    ):
        add_number(doc, text)
    doc.add_heading("Outputs", level=2)
    add_table(
        doc,
        ["Output", "Why it matters"],
        [
            ("State / status word", "Tells the supervisor whether the axis is disabled, ready, running, or faulted"),
            ("Actual position and velocity", "Shows physical/simulated response"),
            ("Following error", "Quantifies reference tracking"),
            ("Fault code", "Makes failure handling explicit and machine-readable"),
            ("Controller heartbeat", "Lets the PLC detect a dead controller"),
            ("Loop statistics", "Mean/max execution time, start jitter, and missed deadlines"),
            ("CSV telemetry", "Supports graphs, comparison, and evidence in the README/demo"),
        ],
        [2800, 6560],
    )

    doc.add_page_break()
    doc.add_heading("4. Parameters and the control cycle", level=1)
    doc.add_heading("Current demonstration parameters", level=2)
    add_table(
        doc,
        ["Group", "Parameter", "Current/default value", "Effect"],
        [
            ("Motion", "Target", "1.57 rad", "Requested final angle"),
            ("Motion", "Max velocity", "0.8 rad/s", "Limits cruise speed"),
            ("Motion", "Max acceleration", "1.5 rad/s²", "Limits speed ramp"),
            ("Completion", "Position tolerance", "0.005 rad", "Position must be this close"),
            ("Completion", "Velocity tolerance", "0.01 rad/s", "Axis must be nearly stopped"),
            ("Watchdog", "PLC timeout", "0.10 s", "Latches communication fault"),
            ("PID", "Kp / Ki / Kd", "32 / 5 / 2", "Tracking response and damping"),
            ("PID", "Derivative filter", "0.008 s", "Reduces derivative noise"),
            ("Actuator", "Torque range", "−10 to +10 N·m", "Saturates controller output"),
            ("Plant", "Inertia / damping", "0.08 / 0.18", "Simulated mechanical response"),
            ("Runtime", "Control frequency", "1,000 Hz", "1 ms nominal cycle"),
            ("Runtime", "PLC worker", "50 Hz", "20 ms supervisory cycle"),
        ],
        [1500, 2200, 2200, 3460],
    )
    doc.add_heading("One 1 ms control cycle", level=2)
    for text in (
        "Read the latest command without blocking.",
        "Check watchdog, E-stop, validity, and controller state.",
        "Advance the trajectory reference.",
        "Calculate and saturate the PID output.",
        "Update the plant and collect feedback.",
        "Write status and timing statistics.",
        "Sleep until the next deadline in real-time-style mode.",
    ):
        add_number(doc, text)
    add_callout(
        doc,
        "Important interpretation",
        "The “1 kHz” value is the requested software cycle. The fast demo skips sleeping, "
        "so its execution-time numbers measure computation rather than Windows scheduling. "
        "A production hard-real-time claim would require an appropriate real-time platform.",
        PALE_AMBER,
    )

    doc.add_page_break()
    doc.add_heading("5. Faults, countermeasures, and recovery", level=1)
    add_table(
        doc,
        ["Fault / hazard", "State", "Detection or countermeasure", "Recovery"],
        [
            ("Emergency stop", "Implemented", "Priority command forces EmergencyStop", "Remove cause, then explicit reset sequence"),
            ("PLC communication timeout · 1001", "Implemented", "Heartbeat must change within 100 ms", "Restore heartbeat, then reset"),
            ("Invalid command · 1006", "Implemented", "Reject NaN, infinity, or invalid limits", "Send valid command, then reset"),
            ("Position limit · 1002", "Defined / next", "Planned soft-position monitor", "Return inside allowed envelope"),
            ("Velocity limit · 1003", "Defined / next", "Planned measured-velocity monitor", "Stop and diagnose command/plant"),
            ("Following error · 1004", "Defined / next", "Planned reference-error threshold", "Diagnose tuning/load, then reset"),
            ("Loop overrun · 1005", "Diagnostic / next", "Deadline-miss counter exists; fault policy next", "Reduce load or move to RT platform"),
            ("Missed short PLC pulse", "Known design risk", "Use held command or sequence/ack handshake", "PLC retains request until acknowledged"),
        ],
        [2050, 1550, 3760, 2000],
        (1, {"Implemented": PALE_GREEN, "Defined / next": PALE_AMBER, "Diagnostic / next": PALE_AMBER, "Known design risk": PALE_RED}),
    )
    doc.add_heading("Fault philosophy", level=2)
    add_bullet(doc, "Faults latch: a momentary recovery does not silently restart motion.")
    add_bullet(doc, "Reset is explicit: the supervisor must request it after the cause is removed.")
    add_bullet(doc, "Safe state dominates: E-stop and active faults override normal motion.")
    add_bullet(doc, "Diagnostics are visible: fault codes and timing counters explain why motion stopped.")
    add_callout(
        doc,
        "Safety boundary",
        "Software safety behavior is demonstrated, but certified machinery safety requires "
        "risk assessment, safety-rated hardware, validated functions, and applicable standards.",
        PALE_RED,
    )

    doc.add_page_break()
    doc.add_heading("6. Evidence and milestone roadmap", level=1)
    doc.add_heading("Measured local evidence", level=2)
    add_table(
        doc,
        ["Result", "Observed value"],
        [
            ("Commanded position", "1.000 rad"),
            ("Final actual position", "1.003 rad"),
            ("Final following error", "−0.003 rad"),
            ("Control cycles to completion", "2,167"),
            ("Deadline misses in reported fast run", "0"),
            ("Asynchronous demo length", "1,500 control cycles"),
            ("Approximate PLC transactions", "75 reads / 69 writes"),
            ("Automated tests", "12 / 12 passing"),
        ],
        [4700, 4660],
    )
    doc.add_paragraph(
        "These values are evidence from the local demonstration, not guaranteed performance "
        "specifications. The generated motion graphs provide the most intuitive visual proof: "
        "the reference accelerates, cruises, decelerates, and the simulated position follows."
    )
    doc.add_heading("Roadmap", level=2)
    add_table(
        doc,
        ["Stage", "Deliverable", "Exit criterion"],
        [
            ("Now", "Local controller + mock PLC + graphs", "Tests pass and motion reaches target"),
            ("Next", "Siemens OPC UA commissioning", "Probe connects and reads exact PLC NodeIds"),
            ("Then", "Live runtime integration", "PLC commands the worker; status returns to PLC"),
            ("Later", "ROS 2 adapter", "Trajectory goal and joint-state telemetry"),
            ("Later", "Fake CiA 402 / EtherCAT", "PDO exchange and drive state machine"),
            ("Stretch", "Native fieldbus / physical drive", "Validated on an appropriate real-time host"),
        ],
        [1200, 3400, 4760],
    )
    doc.add_heading("Tomorrow’s Siemens checklist", level=2)
    for text in (
        "Create the command and status data block in TIA Portal.",
        "Enable the PLC/PLCSIM Advanced OPC UA server.",
        "Expose variables and record the endpoint, namespace, and exact NodeIds.",
        "Copy the example OPC UA configuration and enter those values.",
        "Run the read-only probe; only then connect it to the 50 Hz worker.",
    ):
        add_number(doc, text)

    doc.add_page_break()
    doc.add_heading("7. How to pitch MotionBridge", level=1)
    doc.add_heading("30-second pitch", level=2)
    add_callout(
        doc,
        "Short version",
        "MotionBridge is a modular C++20 industrial motion-control platform. A PLC supervises "
        "enable, motion permission, faults, and heartbeat over OPC UA, while a separate 1 kHz "
        "C++ loop generates a trapezoidal trajectory, runs PID control, and drives a simulated "
        "servo. I separated networking from control through a 50 Hz worker and bounded mailboxes, "
        "so Siemens, ROS 2, and EtherCAT can be added as adapters without changing the tested kernel.",
    )
    doc.add_heading("90-second pitch", level=2)
    doc.add_paragraph(
        "I wanted to demonstrate more than a PID equation, so I built the surrounding industrial "
        "architecture. The axis has a typed state machine, latched faults, emergency-stop behavior, "
        "command validation, heartbeat supervision, timing diagnostics, and tests. It runs fully "
        "locally with a mock PLC and simulated motor, which makes failures repeatable and lets me "
        "show graphs of position, velocity, torque, and following error. The same PLC interface also "
        "has an optional open62541 implementation. The next milestone is to connect it to a Siemens "
        "PLC or PLCSIM Advanced server, then add ROS 2 and a CiA 402-style fieldbus backend. The main "
        "engineering lesson is that OPC UA belongs in the supervisory timing domain—not inside the "
        "1 kHz control loop."
    )
    doc.add_heading("Questions you are likely to get", level=2)
    add_table(
        doc,
        ["Question", "Strong answer"],
        [
            ("Why use a mock PLC?", "It implements the same interface, enables repeatable testing, and proves fault behavior without vendor hardware."),
            ("Why not run OPC UA at 1 kHz?", "Network latency is nondeterministic; OPC UA handles supervisory data while the control loop stays isolated."),
            ("Is this hard real-time?", "No. It is fixed-period, instrumented software on Windows. Hard real-time needs an RT-capable runtime and validation."),
            ("Why simulate the motor?", "It closes the feedback loop and makes controller response, tuning, saturation, and faults visible before hardware."),
            ("How is a fault recovered?", "The fault latches, its cause must disappear, and the PLC must issue an explicit reset."),
            ("What would you improve next?", "Sequence/ack command handshakes, active limit monitors, authenticated OPC UA, then ROS 2 and CiA 402 adapters."),
            ("What was a useful failure?", "A short start pulse could be missed by a slower worker; production logic should hold commands until acknowledged."),
        ],
        [2900, 6460],
    )
    doc.add_heading("Questions you can ask the interviewer", level=2)
    add_bullet(doc, "Where does your architecture draw the boundary between PLC sequencing and motion control?")
    add_bullet(doc, "Which commands use level/ack handshakes rather than one-cycle pulses?")
    add_bullet(doc, "How do you validate loop timing and communication-loss behavior before hardware commissioning?")
    add_bullet(doc, "Which fieldbus and drive profile do you use, and how is the hardware layer simulated in CI?")
    add_bullet(doc, "What safety functions are implemented in certified hardware versus standard control software?")

    doc.add_heading("8. Quick commands", level=1)
    add_callout(
        doc,
        "Build and verify (PowerShell)",
        '& "C:\\Program Files (x86)\\Microsoft Visual Studio\\18\\BuildTools\\Common7\\IDE\\'
        'CommonExtensions\\Microsoft\\CMake\\CMake\\bin\\cmake.exe" --build build\n'
        '& "C:\\Program Files (x86)\\Microsoft Visual Studio\\18\\BuildTools\\Common7\\IDE\\'
        'CommonExtensions\\Microsoft\\CMake\\CMake\\bin\\ctest.exe" --test-dir build --output-on-failure',
        LIGHT_GREY,
    )
    add_callout(
        doc,
        "Run the visual demo",
        r".\build\motionbridge_demo.exe --target 1.0 --max-velocity 0.6 "
        r"--max-acceleration 1.2 --duration 5 --fast --csv build\motion.csv"
        "\n"
        r"python scripts\plot_telemetry.py build\motion.csv --output build\motion.svg",
        LIGHT_GREY,
    )
    add_callout(
        doc,
        "Run the PLC demonstrations",
        r".\build\motionbridge_plc_demo.exe" + "\n" + r".\build\motionbridge_async_plc_demo.exe",
        LIGHT_GREY,
    )
    doc.add_paragraph(
        "The best demonstration sequence is: show the graph, explain the 1 kHz/50 Hz split, "
        "trigger a communication fault, restore the heartbeat, and perform an explicit reset."
    )

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    create_document()
